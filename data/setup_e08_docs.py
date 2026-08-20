"""
E08 document set — builds the multi-format corpus used in the RAG exercise.

Creates three files next to this script:
  store_systems_runbook.docx   Word  — runbook RB-101/RB-102 with a threshold table
  supplier_returns_policy.pdf  PDF   — marketplace supplier returns & chargeback policy
  shipment_status_memo.txt     Text  — a DATED snapshot, deliberately stale (see E08 Part 11)

Run:  python setup_e08_docs.py
"""
import os

HERE = os.path.dirname(os.path.abspath(__file__))


# ── 1. Word document ──────────────────────────────────────────────────
def build_docx():
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading("Store Systems Operations Runbook", level=0)
    p = doc.add_paragraph("Version 7.1  |  Internal Use Only")
    p.add_run("\nOwner: Store Systems Reliability Team").font.size = Pt(10)

    doc.add_heading("RB-101 — POS Transaction Sync Failure", level=1)
    doc.add_heading("Symptom", level=2)
    doc.add_paragraph(
        "Store point-of-sale terminals complete sales locally, but transactions stop "
        "appearing in the central sales ledger. The POS_SYNC_LAG alert fires when the "
        "sync backlog for any store exceeds 750 transactions, or sync delay exceeds 12 minutes."
    )

    doc.add_heading("Alert Thresholds", level=2)
    rows = [
        ("Alert", "Warning", "Critical", "Paging team"),
        ("POS_SYNC_LAG", "750 transactions", "2,000 transactions", "Store Systems Reliability"),
        ("CHECKOUT_LATENCY", "4 seconds", "10 seconds", "Store Systems Reliability"),
        ("INVENTORY_FEED_STALE", "45 minutes", "90 minutes", "Supply Chain Platform"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.cell(r, c)
            cell.text = val
            if r == 0:
                cell.paragraphs[0].runs[0].bold = True

    doc.add_heading("Remediation", level=2)
    for line in [
        "Stuck consumer group: restart consumers ONE AT A TIME. Never restart the whole "
        "group at once — this triggers a rebalance storm and extends the outage.",
        "Bad deploy of the ingestion service: execute the standard one-line rollback, then "
        "let the backlog drain. Expected drain rate is roughly 50,000 transactions per minute.",
        "NEVER manually delete or replay queued POS transactions. Terminals queue locally for "
        "up to 72 hours; manual replay causes duplicate ledger entries and loyalty double-accrual.",
    ]:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("RB-102 — Checkout Latency Degradation", level=1)
    doc.add_paragraph(
        "CHECKOUT_LATENCY fires when median checkout completion exceeds 4 seconds at any "
        "store. Above 10 seconds, declare a SEV-2 and notify the store manager directly. "
        "The most common cause is payment-gateway token cache expiry, not network loss."
    )

    out = os.path.join(HERE, "store_systems_runbook.docx")
    doc.save(out)
    return out


# ── 2. PDF document ───────────────────────────────────────────────────
def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    styles = getSampleStyleSheet()
    out = os.path.join(HERE, "supplier_returns_policy.pdf")
    story = []

    def h(text, level=1):
        story.append(Paragraph(text, styles["Heading%d" % level]))
        story.append(Spacer(1, 6))

    def p(text):
        story.append(Paragraph(text, styles["BodyText"]))
        story.append(Spacer(1, 6))

    h("Marketplace Supplier Returns &amp; Chargeback Policy", 1)
    p("Version 3.4  |  Internal Use Only  |  Owner: Marketplace Seller Operations")

    h("1. Return Window", 2)
    p("Customers may return marketplace items within <b>30 calendar days</b> of delivery. "
      "Electronics carry an extended <b>45-day</b> window. Perishable and clearance items "
      "are final sale and are not eligible for return.")

    h("2. Chargeback Liability", 2)
    p("Where a return results from seller error — wrong item shipped, item not as described, "
      "or damage in seller-controlled packaging — the seller is charged a "
      "<b>flat handling fee of $6.50 per unit</b> plus return shipping. Where the return "
      "results from carrier damage, the platform absorbs the cost and the seller is not charged.")

    h("3. Dispute Escalation", 2)
    p("A seller may dispute a chargeback within <b>14 days</b> of assessment by opening a "
      "case with Seller Operations. Disputes are resolved within 10 business days. "
      "An unresolved dispute automatically escalates to the Marketplace Trust Council "
      "on day 11 — sellers do not need to escalate manually.")

    h("4. Suspension Triggers", 2)
    p("A seller account is placed under review when the seller-fault return rate exceeds "
      "<b>4% of orders</b> over a rolling 60-day window. Two consecutive review periods "
      "above threshold trigger listing suspension. Product safety recalls bypass this "
      "process entirely and result in immediate delisting.")

    SimpleDocTemplate(out, pagesize=A4, title="Marketplace Supplier Returns & Chargeback Policy").build(story)
    return out


# ── 3. The deliberately stale memo ────────────────────────────────────
def build_stale_memo():
    text = """INBOUND SHIPMENT STATUS MEMO
Issued: 10 August 2026, 08:00
Owner: Regional Inbound Logistics — DC 6094
Distribution: Store Systems, Replenishment, Store 4479 management

SUMMARY OF INBOUND SHIPMENTS TO STORE 4479

SHP-88121 — WM-KETTLE-01 — 48 units
  Status: IN TRANSIT, ON SCHEDULE
  Expected arrival: 11 August 2026
  Notes: Trailer loaded at DC 6094 and departed on plan. No exceptions
  reported. Replenishment team should plan shelf reset for the morning
  of 12 August.

SHP-88377 — WM-AIRFRY-11 — 24 units
  Status: IN TRANSIT, ON SCHEDULE
  Expected arrival: 15 August 2026

SHP-88402 — WM-TOWEL-04 — 120 units
  Status: IN TRANSIT
  Expected arrival: 12 August 2026, destination Store 2091

This memo reflects carrier status as of the issue date and time above.
For current status, consult the inbound shipment tracking system.
"""
    out = os.path.join(HERE, "shipment_status_memo.txt")
    with open(out, "w", encoding="utf-8") as f:
        f.write(text)
    return out


if __name__ == "__main__":
    for fn in (build_docx, build_pdf, build_stale_memo):
        path = fn()
        print(f"  created  {os.path.basename(path):32} {os.path.getsize(path):>7,} bytes")
    print("\nE08 document set ready.")
